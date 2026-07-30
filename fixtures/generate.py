"""Generate synthetic fixture documents (fictional data only).

Everything here is invented for testing — no real-world documents are used.
Generated files go to fixtures/generated/ (git-ignored).

Usage:  python fixtures/generate.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Font, PatternFill
from openpyxl.worksheet.datavalidation import DataValidation

OUT = Path(__file__).parent / "generated"

# Same pixel dimensions, different content: swapping one for the other is a
# pure content change, exactly the "logo quietly replaced" case.
LOGO_COLORS = {"logo_blue.png": (0, 90, 200), "logo_red.png": (200, 30, 30)}


def make_logos(out: Path) -> dict[str, Path]:
    """Write the tiny PNGs used by the image fixtures."""
    from PIL import Image as PILImage

    paths = {}
    for name, color in LOGO_COLORS.items():
        path = out / name
        PILImage.new("RGB", (48, 48), color).save(path)
        paths[name] = path
    return paths


def make_excel_baseline(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "計画表"
    ws["A1"] = "架空市 水道整備計画(サンプル)"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:D1")
    headers = ["項目", "数量", "単価", "金額"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDDDDD")
    items = [("管路A", 100, 5000), ("管路B", 250, 3200), ("バルブ", 12, 45000)]
    for i, (name, qty, price) in enumerate(items, start=4):
        ws.cell(row=i, column=1, value=name)
        ws.cell(row=i, column=2, value=qty)
        ws.cell(row=i, column=3, value=price)
        ws.cell(row=i, column=4, value=f"=B{i}*C{i}")
    ws["D8"] = "=SUM(D4:D6)"
    ws["D8"].font = Font(bold=True)

    dv = DataValidation(type="whole", operator="greaterThan", formula1="0")
    dv.add("B4:B6")
    ws.add_data_validation(dv)
    ws.conditional_formatting.add(
        "B4:B6",
        CellIsRule(operator="greaterThan", formula=["200"],
                   fill=PatternFill("solid", fgColor="FFC7CE")),
    )
    ws.oddHeader.center.text = "架空市サンプル資料"
    ws.page_setup.orientation = "landscape"
    ws.print_area = "A1:D8"

    memo = wb.create_sheet("メモ")
    memo["A1"] = "これは架空データです"
    wb.save(path)


def make_excel_good_edit(baseline: Path, path: Path) -> None:
    """Allowed edit: only quantity values in B4:B6 change."""
    from openpyxl import load_workbook

    wb = load_workbook(baseline)
    ws = wb["計画表"]
    ws["B4"] = 120
    ws["B5"] = 300
    wb.save(path)


def make_excel_bad_edit(baseline: Path, path: Path) -> None:
    """Broken edit: overwrites a formula, changes a header format, deletes a sheet."""
    from openpyxl import load_workbook
    from openpyxl.styles import Font

    wb = load_workbook(baseline)
    ws = wb["計画表"]
    ws["B4"] = 120
    ws["D8"] = 9999999          # formula destroyed
    ws["A3"].font = Font(bold=False, color="FF0000")
    del wb["メモ"]
    wb.save(path)


def make_excel_with_image(path: Path, logo: Path | None) -> None:
    """A sheet whose only variable is the embedded logo (or its absence)."""
    wb = Workbook()
    ws = wb.active
    ws.title = "表紙"
    ws["A1"] = "架空市 表紙(サンプル)"
    ws["A3"] = "承認印:"
    if logo is not None:
        ws.add_image(XLImage(str(logo)), "C3")
    wb.save(path)


def make_word_with_image(path: Path, logo: Path | None) -> None:
    doc = Document()
    doc.add_paragraph("架空市 表紙(サンプル)")
    if logo is not None:
        doc.add_picture(str(logo))
    doc.add_paragraph("承認印: 上記のとおり")
    doc.save(path)


def make_word_reformatted(baseline: Path, path: Path) -> None:
    """Text untouched, but rendered invisible: 4pt white.

    No paragraph style changes, so only run-level formatting capture sees it.
    """
    doc = Document(str(baseline))
    for run in doc.paragraphs[3].runs:
        run.font.size = Pt(4)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.save(path)


def make_word_inserted_paragraph(baseline: Path, path: Path) -> None:
    """One paragraph inserted mid-document; every later one shifts down."""
    doc = Document(str(baseline))
    new = doc.add_paragraph("補足: 本節は追記された段落です。")
    doc.paragraphs[1]._p.addnext(new._p)
    doc.save(path)


def make_word_baseline(path: Path) -> None:
    doc = Document()
    doc.add_heading("架空市 業務計画書(サンプル)", level=1)
    doc.add_paragraph("本書は架空のサンプル文書です。")
    doc.add_paragraph("第1章 概要")
    doc.add_paragraph("架空市の水道事業の概要を記載する。")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "項目"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "期間"
    table.cell(1, 1).text = "2026年度〜2030年度"
    table.cell(2, 0).text = "担当"
    table.cell(2, 1).text = "架空部署"
    doc.save(path)


def make_word_good_edit(baseline: Path, path: Path) -> None:
    doc = Document(str(baseline))
    doc.paragraphs[3].text = "架空市の水道事業の概要と更新方針を記載する。"
    doc.save(path)


def make_word_bad_edit(baseline: Path, path: Path) -> None:
    doc = Document(str(baseline))
    doc.paragraphs[3].text = "架空市の水道事業の概要と更新方針を記載する。"
    doc.tables[0].cell(1, 1).text = "2027年度〜2031年度"  # table not allowed
    doc.save(path)


def generate_all(out: Path = OUT) -> dict[str, Path]:
    out.mkdir(parents=True, exist_ok=True)
    paths = {
        "excel_baseline": out / "excel_baseline.xlsx",
        "excel_good": out / "excel_good.xlsx",
        "excel_bad": out / "excel_bad.xlsx",
        "excel_image_baseline": out / "excel_image_baseline.xlsx",
        "excel_image_swapped": out / "excel_image_swapped.xlsx",
        "excel_image_removed": out / "excel_image_removed.xlsx",
        "word_baseline": out / "word_baseline.docx",
        "word_good": out / "word_good.docx",
        "word_bad": out / "word_bad.docx",
        "word_reformatted": out / "word_reformatted.docx",
        "word_inserted": out / "word_inserted.docx",
        "word_image_baseline": out / "word_image_baseline.docx",
        "word_image_swapped": out / "word_image_swapped.docx",
        "word_image_removed": out / "word_image_removed.docx",
    }
    make_excel_baseline(paths["excel_baseline"])
    make_excel_good_edit(paths["excel_baseline"], paths["excel_good"])
    make_excel_bad_edit(paths["excel_baseline"], paths["excel_bad"])
    make_word_baseline(paths["word_baseline"])
    make_word_good_edit(paths["word_baseline"], paths["word_good"])
    make_word_bad_edit(paths["word_baseline"], paths["word_bad"])
    make_word_reformatted(paths["word_baseline"], paths["word_reformatted"])
    make_word_inserted_paragraph(paths["word_baseline"], paths["word_inserted"])

    logos = make_logos(out)
    blue, red = logos["logo_blue.png"], logos["logo_red.png"]
    make_excel_with_image(paths["excel_image_baseline"], blue)
    make_excel_with_image(paths["excel_image_swapped"], red)
    make_excel_with_image(paths["excel_image_removed"], None)
    make_word_with_image(paths["word_image_baseline"], blue)
    make_word_with_image(paths["word_image_swapped"], red)
    make_word_with_image(paths["word_image_removed"], None)
    return paths


if __name__ == "__main__":
    for name, p in generate_all().items():
        print(f"{name}: {p}")
