"""Reports have to be readable by whoever approves the change.

A violation that says "None -> None" or lists fifty-six shifted paragraphs
tells a reviewer nothing they can act on, so these tests pin what the report
actually says, not just that something was reported.
"""

import json
import zipfile

import pytest
from docx import Document
from openpyxl import Workbook

from templategate import check, diff
from templategate.core.policy import parse_policy
from templategate.core.selector import match_selector, quote_sheet
from templategate.reporters import render_json, render_markdown, render_text

STRICT_WORD = parse_policy({
    "target": "word",
    "protect": [{"selector": "*", "attributes": ["*"]}],
})


def _only(changes, attribute):
    return [c for c in changes if c.attribute == attribute]


# --- C1: a format change says which property changed ---------------------

def test_number_format_change_names_the_property(tmp_path):
    baseline = tmp_path / "base.xlsx"
    candidate = tmp_path / "cand.xlsx"
    for path, number_format in ((baseline, "#,##0"), (candidate, "#,##0,")):
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A4"] = 1000
        ws["A4"].number_format = number_format
        wb.save(path)

    change = _only(diff(baseline, candidate), "format")[0]
    assert change.detail == "cell format changed: numfmt #,##0 -> #,##0,"
    assert change.old == {"numfmt": "#,##0"}
    assert change.new == {"numfmt": "#,##0,"}


def test_format_delta_carries_only_the_changed_property(tmp_path):
    """The reader wants the one thing that moved, not the whole style."""
    from openpyxl.styles import Font

    baseline = tmp_path / "base.xlsx"
    candidate = tmp_path / "cand.xlsx"
    for path, color in ((baseline, "FF000000"), (candidate, "FFFFFFFF")):
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "total"
        ws["A1"].font = Font(bold=True, color=color)
        wb.save(path)

    change = _only(diff(baseline, candidate), "format")[0]
    assert list(change.old) == ["font.color"]
    assert "font.color" in change.detail


def test_word_hidden_text_detail_names_the_property(tmp_path):
    from generate import rewrite_zip

    def build(name, run_properties):
        plain = tmp_path / f"{name}_plain.docx"
        document = Document()
        document.add_paragraph("x")
        document.save(str(plain))
        with zipfile.ZipFile(plain) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        body = (f"<w:p><w:r>{run_properties}<w:t>Liability is capped.</w:t>"
                "</w:r></w:p>")
        built = tmp_path / f"{name}.docx"
        rewrite_zip(plain, built, add={
            "word/document.xml": xml.replace("</w:body>", body + "</w:body>").encode()})
        return built

    baseline = build("base", "")
    candidate = build("cand", "<w:rPr><w:vanish/></w:rPr>")
    change = _only(diff(baseline, candidate), "format")[0]
    assert "vanish" in change.detail


# --- C2: a sheet name containing "!" is not the same as a cell reference --

def test_quote_sheet_only_quotes_what_needs_it():
    assert quote_sheet("Sheet1") == "Sheet1"
    assert quote_sheet("計画表") == "計画表"
    assert quote_sheet("Q1!Q4") == "'Q1!Q4'"
    assert quote_sheet("Q1#Draft") == "'Q1#Draft'"
    assert quote_sheet("It's") == "'It''s'"


def test_cell_selector_never_reaches_a_same_named_sheet():
    """"Q1!Q4" is cell Q4 of sheet Q1, and must stay only that."""
    assert match_selector("Q1!Q4", "Q1!Q4")
    assert not match_selector("Q1!Q4", "'Q1!Q4'!A1")
    assert not match_selector("Q1!Q4", "'Q1!Q4'!Q4")


def test_quoted_sheet_selector_never_reaches_the_cell():
    assert match_selector("'Q1!Q4'!A1", "'Q1!Q4'!A1")
    assert match_selector("'Q1!Q4'", "'Q1!Q4'!A1")
    assert not match_selector("'Q1!Q4'", "Q1!Q4")
    assert not match_selector("'Q1!Q4'!A1:B9", "Q1!Q4")


def test_locations_quote_ambiguous_sheet_names(tmp_path):
    """A workbook with both readings present must keep them apart."""
    def build(path, secret):
        wb = Workbook()
        ws = wb.active
        ws.title = "Q1"
        ws["Q4"] = secret
        wb.create_sheet("Q1!Q4")["A1"] = "scratch"
        wb.save(path)

    baseline = tmp_path / "base.xlsx"
    candidate = tmp_path / "cand.xlsx"
    build(baseline, "PROTECTED secret")
    build(candidate, "LEAKED")

    locations = {c.location for c in diff(baseline, candidate)}
    assert locations == {"Q1!Q4"}

    # A rule written for the scratch *sheet* must not touch the secret cell.
    policy = parse_policy({
        "target": "excel",
        "allow": [{"selector": "'Q1!Q4'", "attributes": ["value"]}],
    })
    assert not check(baseline, candidate, policy).passed


def test_edits_on_the_ambiguous_sheet_are_addressable(tmp_path):
    def build(path, scratch):
        wb = Workbook()
        ws = wb.active
        ws.title = "Q1"
        ws["Q4"] = "secret"
        wb.create_sheet("Q1!Q4")["A1"] = scratch
        wb.save(path)

    baseline = tmp_path / "base.xlsx"
    candidate = tmp_path / "cand.xlsx"
    build(baseline, "scratch")
    build(candidate, "rewritten")

    assert {c.location for c in diff(baseline, candidate)} == {"'Q1!Q4'!A1"}
    policy = parse_policy({
        "target": "excel",
        "allow": [{"selector": "'Q1!Q4'", "attributes": ["value"]}],
    })
    assert check(baseline, candidate, policy).passed


# --- C3: one edit reads as one edit --------------------------------------

def _numbered(tmp_path, name, skip=None):
    path = tmp_path / f"{name}.docx"
    document = Document()
    for number in range(1, 21):
        if number == skip:
            continue
        document.add_paragraph(f"Clause {number}.")
    document.save(str(path))
    return path


def test_one_deletion_still_reports_every_change(tmp_path):
    """Collapsing is a reporting choice; the policy still sees everything."""
    baseline = _numbered(tmp_path, "base")
    candidate = _numbered(tmp_path, "cand", skip=5)
    changes = diff(baseline, candidate)
    assert len(changes) > 10
    result = check(baseline, candidate, STRICT_WORD)
    assert not result.passed
    assert len(result.violations) == len(changes)


def test_shifted_paragraphs_are_grouped_under_the_real_edit(tmp_path):
    baseline = _numbered(tmp_path, "base")
    candidate = _numbered(tmp_path, "cand", skip=5)
    groups = {c.group for c in diff(baseline, candidate) if c.group}
    assert groups == {"1 paragraph removed at p5"}


def test_text_report_collapses_the_cascade(tmp_path):
    baseline = _numbered(tmp_path, "base")
    candidate = _numbered(tmp_path, "cand", skip=5)
    report = render_text(check(baseline, candidate, STRICT_WORD))
    violation_lines = [line for line in report.splitlines()
                       if line.startswith("  [error]")]
    assert len(violation_lines) <= 4, report
    assert any("content shifted because 1 paragraph removed at p5" in line
               for line in report.splitlines())


def test_json_report_keeps_every_change(tmp_path):
    """Machines get the full list even when humans get the summary."""
    baseline = _numbered(tmp_path, "base")
    candidate = _numbered(tmp_path, "cand", skip=5)
    result = check(baseline, candidate, STRICT_WORD)
    data = json.loads(render_json(result))
    assert len(data["violations"]) == len(result.violations)
    assert any(v["change"]["group"] for v in data["violations"])


def test_markdown_report_collapses_the_cascade(tmp_path):
    baseline = _numbered(tmp_path, "base")
    candidate = _numbered(tmp_path, "cand", skip=5)
    report = render_markdown(check(baseline, candidate, STRICT_WORD))
    rows = [line for line in report.splitlines()
            if line.startswith("| error")]
    assert len(rows) <= 4, report


def _reordered(tmp_path, name, order):
    path = tmp_path / f"{name}.docx"
    document = Document()
    document.add_paragraph("Obligations")
    for text in order:
        document.add_paragraph(text)
    document.save(str(path))
    return path


CLAUSES = ["1. The Buyer shall pay within 14 days.",
           "2. The Seller shall deliver on time.",
           "3. The Supplier may suspend delivery while payment is overdue."]


def test_reordering_is_reported_as_a_move(tmp_path):
    baseline = _reordered(tmp_path, "base", CLAUSES)
    candidate = _reordered(tmp_path, "cand", [CLAUSES[2], CLAUSES[1], CLAUSES[0]])
    moves = _only(diff(baseline, candidate, align=True), "moved")
    assert moves
    assert all("moved from" in m.detail for m in moves)
    assert all(m.old != m.new for m in moves)


def test_page_extension_no_longer_waves_a_reorder_through(tmp_path):
    """Swapping two clauses is not a text edit, so a text rule cannot allow it."""
    baseline = _reordered(tmp_path, "base", CLAUSES)
    candidate = _reordered(tmp_path, "cand", [CLAUSES[2], CLAUSES[1], CLAUSES[0]])
    policy = parse_policy({
        "target": "word",
        "mode": "page_extension",
        "allow": [{"selector": "body", "attributes": ["text"]}],
    })
    assert not check(baseline, candidate, policy).passed


def test_page_extension_still_allows_a_plain_insertion(tmp_path):
    """The relaxation has to keep working: an insertion is not a move."""
    baseline = _reordered(tmp_path, "base", CLAUSES)
    candidate = _reordered(tmp_path, "cand",
                           [CLAUSES[0], "1a. Interest accrues after 14 days.",
                            CLAUSES[1], CLAUSES[2]])
    changes = diff(baseline, candidate, align=True)
    assert not _only(changes, "moved")
    assert [c.detail for c in changes] == ["paragraph added"]


# --- C4: wording a reader can trust --------------------------------------

def _one(change):
    from templategate.core.model import CheckResult, Violation

    return CheckResult(
        passed=False, target="word", baseline="b.docx", candidate="c.docx",
        changes=[change],
        violations=[Violation(change=change, rule="protected",
                              message="protected attribute "
                                      f"{change.attribute!r} changed")])


def test_a_change_the_detail_explains_prints_no_old_and_new(tmp_path):
    """"present -> present" fills two columns and says nothing."""
    from templategate.core.model import Change

    change = Change("section1#footer:default", "header_footer",
                    old=None, new=None,
                    detail="the footer on the default page changed")
    report = render_text(_one(change))
    assert "the footer on the default page changed" in report
    assert "->" not in report.split("Violations:")[1]


def test_the_markdown_row_carries_the_detail_when_old_and_new_do_not(tmp_path):
    from templategate.core.model import Change

    change = Change("section1#footer:default", "header_footer",
                    old=None, new=None,
                    detail="the footer on the default page changed")
    row = [line for line in render_markdown(_one(change)).splitlines()
           if line.startswith("| error")][0]
    assert "the footer on the default page changed" in row
    assert "none" not in row


def test_values_are_still_shown_when_nothing_else_explains_them(tmp_path):
    """Suppression follows the detail, not the shape of old/new.

    The header/footer text summary carries dicts and no detail; hiding those
    would leave the violation with nothing to say at all.
    """
    from templategate.core.model import Change

    change = Change("section1#header_footer", "header_footer",
                    old={"header": "CONFIDENTIAL"}, new={"header": "PUBLIC"})
    report = render_text(_one(change))
    assert "CONFIDENTIAL" in report and "PUBLIC" in report


def test_a_real_removal_still_shows_both_sides(tmp_path):
    from templategate.core.model import Change

    change = Change("section1#footer:default", "header_footer",
                    old="present", new=None,
                    detail="the footer on the default page is gone from this section")
    assert "present -> none" in render_text(_one(change))


def test_no_report_ever_says_present_to_present(fixtures, tmp_path):
    """The whole point: an unchanged-looking pair is never printed."""
    from generate import rewrite_zip

    baseline = tmp_path / "b.docx"
    rewrite_zip(fixtures["word_baseline"], baseline, add={})
    header = ('<?xml version="1.0"?><w:hdr xmlns:w="http://schemas.'
              'openxmlformats.org/wordprocessingml/2006/main"><w:p><w:r>'
              "<w:t>{}</w:t></w:r></w:p></w:hdr>")
    for report in (render_text, render_markdown):
        rendered = report(check(baseline, baseline, STRICT_WORD))
        assert "present -> present" not in rendered
        assert "| present | present |" not in rendered


def test_a_nudged_image_reads_as_one_move(tmp_path):
    """The same picture at a new anchor is a move, not a replacement.

    It arrives as a removal and an addition carrying identical content
    hashes; printing both reads as if the picture might be gone.
    """
    from openpyxl.drawing.image import Image as XLImage
    from PIL import Image as PILImage

    logo = tmp_path / "logo.png"
    PILImage.new("RGB", (40, 40), (10, 90, 200)).save(logo)

    def build(path, anchor):
        wb = Workbook()
        ws = wb.active
        ws.title = "見積書"
        ws["A1"] = "x"
        ws.add_image(XLImage(str(logo)), anchor)
        wb.save(path)
        return path

    changes = diff(build(tmp_path / "b.xlsx", "C3"),
                   build(tmp_path / "c.xlsx", "H12"))
    assert len(changes) == 1
    assert changes[0].detail == "same image, moved from C3 to H12"


def test_a_replaced_image_still_reads_as_two_lines(tmp_path):
    """Different content is a replacement, and must not be called a move."""
    from openpyxl.drawing.image import Image as XLImage
    from PIL import Image as PILImage

    def build(path, colour):
        logo = tmp_path / f"logo{colour[0]}.png"
        PILImage.new("RGB", (40, 40), colour).save(logo)
        wb = Workbook()
        ws = wb.active
        ws.title = "見積書"
        ws["A1"] = "x"
        ws.add_image(XLImage(str(logo)), "C3")
        wb.save(path)
        return path

    changes = diff(build(tmp_path / "b.xlsx", (10, 90, 200)),
                   build(tmp_path / "c.xlsx", (200, 30, 30)))
    assert len(changes) == 2
    assert {c.detail for c in changes} == {"image removed from C3",
                                           "image added at C3"}


def test_a_resized_image_says_resized():
    """Exercised directly: openpyxl re-derives size from the file on load."""
    from templategate.excel.diff import _diff_images

    sha = "a" * 64
    changes = _diff_images("Sheet1", {(sha, (2, 2), (40, 40))},
                           {(sha, (2, 2), (120, 120))})
    assert [c.detail for c in changes] == [
        "same image, resized from 40x40 to 120x120"]


def test_an_image_both_moved_and_resized_says_both():
    from templategate.excel.diff import _diff_images

    sha = "b" * 64
    changes = _diff_images("Sheet1", {(sha, (2, 2), (40, 40))},
                           {(sha, (7, 11), (120, 120))})
    assert changes[0].detail == (
        "same image, moved from C3 to H12 and resized from 40x40 to 120x120")


def test_a_layout_change_reads_in_words(tmp_path):
    def build(path, width, hidden):
        wb = Workbook()
        ws = wb.active
        ws.title = "見積書"
        ws["A1"] = "x"
        ws.column_dimensions["H"].width = width
        ws.column_dimensions["H"].hidden = hidden
        wb.save(path)
        return path

    widened = diff(build(tmp_path / "b.xlsx", 8.43, True),
                   build(tmp_path / "c.xlsx", 13, True))
    assert widened[0].detail == (
        "row or column changed: column width 8.43 -> 13.0 (still hidden)")

    unhidden = diff(build(tmp_path / "d.xlsx", 8.43, True),
                    build(tmp_path / "e.xlsx", 8.43, False))
    assert unhidden[0].detail == "row or column changed: hidden True -> False"


def test_a_row_height_change_says_row_height(tmp_path):
    def build(path, height):
        wb = Workbook()
        ws = wb.active
        ws.title = "見積書"
        ws["A1"] = "x"
        ws.row_dimensions[4].height = height
        wb.save(path)
        return path

    changes = diff(build(tmp_path / "b.xlsx", 15), build(tmp_path / "c.xlsx", 40))
    assert changes[0].detail == "row or column changed: row height 15.0 -> 40.0"


def test_merged_cells_say_so(tmp_path):
    def build(path, merge):
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"], ws["B1"] = "Q1", "Q2"
        if merge:
            ws.merge_cells("A1:B1")
        wb.save(path)

    baseline = tmp_path / "base.xlsx"
    candidate = tmp_path / "cand.xlsx"
    build(baseline, False)
    build(candidate, True)
    merged = _only(diff(baseline, candidate), "merge")
    assert merged and "cells merged" in merged[0].detail

    unmerged = _only(diff(candidate, baseline), "merge")
    assert unmerged and "cells unmerged" in unmerged[0].detail


def test_macro_enabled_workbook_saved_as_xlsx_is_reported(tmp_path):
    from generate import rewrite_zip

    baseline = tmp_path / "base.xlsm"
    candidate = tmp_path / "cand.xlsx"
    wb = Workbook()
    wb.active["A1"] = "x"
    wb.save(baseline)
    rewrite_zip(baseline, candidate, add={})

    change = [c for c in diff(baseline, candidate)
              if c.location == "workbook#format"][0]
    assert (change.old, change.new) == ("xlsm", "xlsx")
    assert "loses its VBA project" in change.detail


def test_page_break_removal_says_page_break(tmp_path):
    from generate import rewrite_zip

    def build(name, body):
        plain = tmp_path / f"{name}_plain.docx"
        document = Document()
        document.add_paragraph("x")
        document.save(str(plain))
        with zipfile.ZipFile(plain) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        built = tmp_path / f"{name}.docx"
        rewrite_zip(plain, built, add={
            "word/document.xml": xml.replace("</w:body>", body + "</w:body>").encode()})
        return built

    baseline = build("base", '<w:p><w:r><w:br w:type="page"/></w:r></w:p>')
    candidate = build("cand", "<w:p><w:r></w:r></w:p>")
    breaks = [c for c in diff(baseline, candidate)
              if "page break" in c.detail]
    assert breaks
    assert "page break removed" in breaks[0].detail
