"""Word rewriting its own file is not an edit.

Opening a document in Word and pressing save changes a surprising amount of
the package: session ids are restamped, the shape-id allocator advances, the
font table gains entries for the locale, unused header parts are discarded
and the rest renumbered.  None of it changes a word on the page, and all of
it was reported as damage until each class was pinned here.

Every quietening is paired with a real change of the same shape, because a
normalization that also swallows the attack has made things worse.
"""

import zipfile

import pytest
from docx import Document

from templategate import check, diff, snapshot
from templategate.core.package import take_package_snapshot
from templategate.core.policy import parse_policy

STRICT = parse_policy({
    "target": "word",
    "protect": [{"selector": "*", "attributes": ["*"]}],
})


def _document(path, *, header=None, paragraphs=("Body text.",)):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if header is not None:
        document.sections[0].header.paragraphs[0].text = header
    document.save(str(path))
    return path


def _rewrite(source, destination, *, replace=None, rename=None):
    """Rewrite a package the way another save would: patch and/or renumber.

    A renamed part has to be renamed everywhere it is referred to, or the
    document simply will not open — which is a different test.
    """
    replace = dict(replace or {})
    rename = rename or {}
    if rename:
        def renumber(xml: str) -> str:
            for old, new in rename.items():
                xml = xml.replace("/" + old, "/" + new)
            return xml

        previous = replace.get("[Content_Types].xml")
        replace["[Content_Types].xml"] = (
            (lambda xml: renumber(previous(xml))) if previous else renumber)

    with zipfile.ZipFile(source) as zin:
        items = [(item.filename, zin.read(item.filename)) for item in zin.infolist()]
    with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items:
            if name in replace:
                data = replace[name](data.decode("utf-8")).encode("utf-8")
            zout.writestr(rename.get(name, name), data)
    return destination


def _attributes(changes):
    return {c.attribute for c in changes}


# --- session bookkeeping --------------------------------------------------

def test_revision_session_ids_are_not_a_change(tmp_path):
    """Word restamps w:rsid elements and paraId/textId on every save."""
    baseline = _document(tmp_path / "b.docx")

    def restamp(xml: str) -> str:
        xml = xml.replace("<w:rsids>",
                          '<w:rsids><w:rsid w:val="00AB12CD"/>'
                          '<w:rsid w:val="00FF3344"/>')
        return xml

    def restamp_ids(xml: str) -> str:
        return xml.replace("<w:p>", '<w:p w14:paraId="11112222" '
                           'w14:textId="33334444" '
                           'xmlns:w14="http://schemas.microsoft.com/office/word'
                           '/2010/wordml">')

    candidate = _rewrite(baseline, tmp_path / "c.docx", replace={
        "word/settings.xml": restamp,
        "word/document.xml": restamp_ids,
    })
    assert diff(baseline, candidate) == []


def test_the_shape_id_allocator_is_not_a_change(tmp_path):
    """Word bumps spidmax and adds header shape defaults as it goes."""
    baseline = _document(tmp_path / "b.docx")

    def advance(xml: str) -> str:
        return xml.replace(
            "</w:settings>",
            '<w:hdrShapeDefaults><o:shapedefaults v:ext="edit" spidmax="2050" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:v="urn:schemas-microsoft-com:vml"/></w:hdrShapeDefaults>'
            "</w:settings>")

    candidate = _rewrite(baseline, tmp_path / "c.docx",
                         replace={"word/settings.xml": advance})
    assert diff(baseline, candidate) == []


def test_document_protection_is_still_compared(tmp_path):
    """settings.xml is quietened, not ignored: locking still has to show."""
    baseline = _document(tmp_path / "b.docx")
    candidate = _rewrite(baseline, tmp_path / "c.docx", replace={
        "word/settings.xml": lambda xml: xml.replace(
            "</w:settings>",
            '<w:documentProtection w:edit="readOnly" w:enforcement="1"/>'
            "</w:settings>")})
    assert _attributes(diff(baseline, candidate)) == {"parts"}


# --- the font table is a manifest, not content ---------------------------

def test_font_table_entries_are_not_content(tmp_path):
    """Word adds the fonts its locale needs; the text still uses what it used."""
    baseline = _document(tmp_path / "b.docx")
    candidate = _rewrite(baseline, tmp_path / "c.docx", replace={
        "word/fontTable.xml": lambda xml: xml.replace(
            "</w:fonts>",
            '<w:font w:name="游明朝"><w:charset w:val="80"/></w:font></w:fonts>')})
    assert diff(baseline, candidate) == []


def test_a_font_actually_applied_to_text_still_fails(tmp_path):
    """The font table is quiet; the font on a run is not."""
    baseline = _document(tmp_path / "b.docx")
    candidate = _rewrite(baseline, tmp_path / "c.docx", replace={
        "word/document.xml": lambda xml: xml.replace(
            "<w:r>", '<w:r><w:rPr><w:rFonts w:ascii="Papyrus"/></w:rPr>', 1)})
    assert "format" in _attributes(diff(baseline, candidate))


# --- generated furniture --------------------------------------------------

def test_regenerated_footnote_separators_are_not_content(tmp_path):
    """Word rewrites the rule drawn above footnotes on every save."""
    baseline = _document(tmp_path / "b.docx")
    footnotes = (
        '<?xml version="1.0"?><w:footnotes xmlns:w="http://schemas.'
        'openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r>'
        "<w:separator/></w:r></w:p></w:footnote>"
        '<w:footnote w:id="1"><w:p><w:r><w:t>Material disclosure.</w:t>'
        "</w:r></w:p></w:footnote></w:footnotes>"
    )
    with_separator = _rewrite(baseline, tmp_path / "b2.docx")
    from generate import rewrite_zip

    rewrite_zip(baseline, with_separator, add={"word/footnotes.xml": footnotes.encode()})
    regenerated = tmp_path / "c.docx"
    rewrite_zip(with_separator, regenerated, add={
        "word/footnotes.xml": footnotes.replace(
            "<w:separator/>",
            '<w:rPr><w:rFonts w:hint="eastAsia"/></w:rPr><w:separator/>').encode()})
    assert diff(with_separator, regenerated) == []


def test_a_real_footnote_change_still_fails(tmp_path):
    from generate import rewrite_zip

    def footnotes(text):
        return (
            '<?xml version="1.0"?><w:footnotes xmlns:w="http://schemas.'
            'openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r>'
            "<w:separator/></w:r></w:p></w:footnote>"
            f'<w:footnote w:id="1"><w:p><w:r><w:t>{text}</w:t>'
            "</w:r></w:p></w:footnote></w:footnotes>"
        ).encode()

    baseline = _document(tmp_path / "b.docx")
    with_note = tmp_path / "b2.docx"
    rewrite_zip(baseline, with_note, add={"word/footnotes.xml": footnotes("Material.")})
    altered = tmp_path / "c.docx"
    rewrite_zip(with_note, altered, add={"word/footnotes.xml": footnotes("Deleted.")})
    assert _attributes(diff(with_note, altered)) == {"parts"}


# --- headers are compared by role ----------------------------------------

def test_header_parts_are_not_compared_by_filename(fixtures, tmp_path):
    document = _document(tmp_path / "b.docx", header="CONFIDENTIAL")
    inventory = take_package_snapshot(document)["parts"]
    assert not [name for name in inventory if name.startswith("word/header")]
    assert "section1#header:default" in snapshot(document)["header_footer_parts"]


def test_renumbering_a_header_is_not_a_change(tmp_path):
    """Word discards unused header parts and renumbers what is left."""
    baseline = _document(tmp_path / "b.docx", header="CONFIDENTIAL")
    with zipfile.ZipFile(baseline) as zf:
        assert "word/header1.xml" in zf.namelist()

    candidate = _rewrite(
        baseline, tmp_path / "c.docx",
        rename={"word/header1.xml": "word/header2.xml"},
        replace={"word/_rels/document.xml.rels":
                 lambda xml: xml.replace("header1.xml", "header2.xml")})
    with zipfile.ZipFile(candidate) as zf:
        assert "word/header2.xml" in zf.namelist()
    assert diff(baseline, candidate) == []


def test_a_renumbered_header_whose_content_changed_still_fails(tmp_path):
    """The renumbering is forgiven; what the header says is not."""
    baseline = _document(tmp_path / "b.docx", header="CONFIDENTIAL")
    candidate = _rewrite(
        baseline, tmp_path / "c.docx",
        rename={"word/header1.xml": "word/header2.xml"},
        replace={
            "word/_rels/document.xml.rels":
                lambda xml: xml.replace("header1.xml", "header2.xml"),
            "word/header1.xml":
                lambda xml: xml.replace("CONFIDENTIAL", "Public"),
        })
    result = check(baseline, candidate, STRICT)
    assert not result.passed
    assert "header_footer" in _attributes(diff(baseline, candidate))


def test_an_unused_header_role_is_not_reported_when_word_drops_it(tmp_path):
    """An even-page header that no setting displays is not on the page."""
    baseline = _document(tmp_path / "b.docx", header="CONFIDENTIAL")
    roles = snapshot(baseline)["header_footer_parts"]
    assert set(roles) == {"section1#header:default"}


def test_a_first_page_header_counts_once_the_document_uses_one(tmp_path):
    path = tmp_path / "b.docx"
    document = Document()
    document.add_paragraph("Body.")
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = "COVER"
    section.header.paragraphs[0].text = "CONFIDENTIAL"
    document.save(str(path))
    assert "section1#header:first" in snapshot(path)["header_footer_parts"]


# --- and the steady state as a whole -------------------------------------

def test_a_resave_that_changes_nothing_reports_nothing(tmp_path):
    """All of the above at once, which is what one Word save looks like."""
    baseline = _document(tmp_path / "b.docx", header="CONFIDENTIAL",
                         paragraphs=("First.", "Second."))
    candidate = _rewrite(
        baseline, tmp_path / "c.docx",
        rename={"word/header1.xml": "word/header3.xml"},
        replace={
            "word/_rels/document.xml.rels":
                lambda xml: xml.replace("header1.xml", "header3.xml"),
            "word/settings.xml":
                lambda xml: xml.replace(
                    "<w:rsids>", '<w:rsids><w:rsid w:val="00AB12CD"/>'),
            "word/fontTable.xml":
                lambda xml: xml.replace(
                    "</w:fonts>",
                    '<w:font w:name="游明朝"><w:charset w:val="80"/></w:font>'
                    "</w:fonts>"),
            "word/document.xml":
                lambda xml: xml.replace(
                    "<w:p>", '<w:p w14:paraId="11112222" '
                    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'),
        })
    assert diff(baseline, candidate) == []
    assert check(baseline, candidate, STRICT).passed
