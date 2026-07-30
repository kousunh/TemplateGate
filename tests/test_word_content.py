"""Content python-docx does not model, read from document.xml directly.

Every case here is a real edit that left ``Document.paragraphs`` and
``Document.tables`` looking identical: text inside a content control or a
text box, a nested table, a field instruction, an edit recorded as a tracked
change, hidden text, a bookmark, a moved tab stop.
"""

import zipfile

import pytest
from docx import Document

from templategate import check, diff, snapshot
from templategate.core.policy import parse_policy

STRICT = parse_policy({
    "target": "word",
    "protect": [{"selector": "*", "attributes": ["*"]}],
})
# What a sane policy looks like: body text may be rewritten, nothing else.
BODY_TEXT_ONLY = parse_policy({
    "target": "word",
    "allow": [{"selector": "body", "attributes": ["text"]}],
    "protect": [{"selector": "*", "attributes": ["style", "format", "section",
                                                 "header_footer", "table"]}],
})

V_NS = 'xmlns:v="urn:schemas-microsoft-com:vml"'


def _build(tmp_path, name, body_xml, paragraphs=("Intro paragraph.",)):
    """A .docx whose body ends with the given raw XML."""
    from generate import rewrite_zip

    plain = tmp_path / f"{name}_plain.docx"
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(plain))

    with zipfile.ZipFile(plain) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    injected = xml.replace("</w:body>", body_xml + "</w:body>")
    built = tmp_path / f"{name}.docx"
    rewrite_zip(plain, built, add={"word/document.xml": injected.encode("utf-8")})
    return built


def _pair(tmp_path, body_before, body_after, **kwargs):
    return (_build(tmp_path, "base", body_before, **kwargs),
            _build(tmp_path, "cand", body_after, **kwargs))


def _by_attribute(changes):
    out = {}
    for change in changes:
        out.setdefault(change.attribute, set()).add(change.location)
    return out


# --- containers python-docx skips ---------------------------------------

def _sdt(text, lock="sdtContentLocked"):
    lock_xml = f'<w:lock w:val="{lock}"/>' if lock else ""
    return (
        f'<w:sdt><w:sdtPr><w:tag w:val="amount"/>{lock_xml}</w:sdtPr>'
        f"<w:sdtContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:sdtContent></w:sdt>"
    )


def test_content_control_text_is_compared(tmp_path):
    baseline, candidate = _pair(tmp_path, _sdt("250,000"), _sdt("25,000"))
    assert _by_attribute(diff(baseline, candidate))["text"] == {"sdt1"}
    assert not check(baseline, candidate, STRICT).passed


def test_content_control_is_not_covered_by_a_body_text_rule(tmp_path):
    """"body" means the body's own paragraphs, not everything nested in it."""
    baseline, candidate = _pair(tmp_path, _sdt("250,000"), _sdt("25,000"))
    assert not check(baseline, candidate, BODY_TEXT_ONLY).passed


def test_content_control_lock_removal_is_compared(tmp_path):
    baseline, candidate = _pair(tmp_path, _sdt("250,000"),
                                _sdt("250,000", lock=None))
    assert _by_attribute(diff(baseline, candidate))["content_control"] == {"sdt1"}


def _textbox(text):
    return (
        f'<w:p><w:r><w:pict><v:shape {V_NS}><v:textbox><w:txbxContent>'
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )


def test_text_box_content_is_compared(tmp_path):
    baseline, candidate = _pair(tmp_path, _textbox("DO NOT operate above 40C"),
                                _textbox("Safe to operate above 40C"))
    assert _by_attribute(diff(baseline, candidate))["text"] == {"textbox1"}
    assert not check(baseline, candidate, BODY_TEXT_ONLY).passed


def _nested_table(rate):
    return (
        "<w:tbl><w:tr><w:tc><w:p><w:r><w:t>outer</w:t></w:r></w:p>"
        f"<w:tbl><w:tr><w:tc><w:p><w:r><w:t>{rate}</w:t></w:r></w:p></w:tc>"
        "</w:tr></w:tbl></w:tc></w:tr></w:tbl>"
    )


def test_nested_table_cell_is_compared(tmp_path):
    baseline, candidate = _pair(tmp_path, _nested_table("USD 300/h"),
                                _nested_table("USD 30/h"))
    assert _by_attribute(diff(baseline, candidate))["text"] == {
        "table1!r1c1!table1!r1c1"
    }
    assert not check(baseline, candidate, BODY_TEXT_ONLY).passed


# --- text that is not where python-docx looks ---------------------------

def _field(instruction, cached):
    return (f'<w:p><w:fldSimple w:instr="{instruction}">'
            f"<w:r><w:t>{cached}</w:t></w:r></w:fldSimple></w:p>")


def test_field_instruction_is_compared(tmp_path):
    """The cached result is ordinary text; the instruction behind it is not."""
    baseline, candidate = _pair(
        tmp_path,
        _field(" DOCPROPERTY EffectiveDate ", "2026-01-01"),
        _field(" DOCPROPERTY SupersededDate ", "2026-01-01"),
    )
    changes = diff(baseline, candidate)
    assert _by_attribute(changes)["field"] == {"p2"}
    assert "text" not in _by_attribute(changes)


def test_tracked_insertion_counts_as_displayed_text(tmp_path):
    """A tracked insertion is on the page, so it is a text change."""
    plain = "<w:p><w:r><w:t>The Supplier is liable.</w:t></w:r></w:p>"
    inserted = (
        "<w:p><w:r><w:t>The Supplier is liable.</w:t></w:r>"
        '<w:ins w:id="9" w:author="a" w:date="2026-01-01T00:00:00Z">'
        "<w:r><w:t> Except where it is not.</w:t></w:r></w:ins></w:p>"
    )
    baseline, candidate = _pair(tmp_path, plain, inserted)
    changes = _by_attribute(diff(baseline, candidate))
    assert changes["text"] == {"p2"}
    assert changes["revision"] == {"p2"}


def test_tracked_deletion_removes_displayed_text(tmp_path):
    """A tracked deletion is struck from the page, so it is a text change."""
    plain = "<w:p><w:r><w:t>Indemnity applies.</w:t></w:r></w:p>"
    deleted = ('<w:p><w:del w:id="9" w:author="a" w:date="2026-01-01T00:00:00Z">'
               "<w:r><w:delText>Indemnity applies.</w:delText></w:r></w:del></w:p>")
    baseline, candidate = _pair(tmp_path, plain, deleted)
    changes = diff(baseline, candidate)
    text = [c for c in changes if c.attribute == "text"]
    assert [c.location for c in text] == ["p2"]
    assert text[0].new == ""


def test_hidden_text_is_reported_as_a_format_change(tmp_path):
    visible = "<w:p><w:r><w:t>Liability is capped.</w:t></w:r></w:p>"
    hidden = ("<w:p><w:r><w:rPr><w:vanish/></w:rPr>"
              "<w:t>Liability is capped.</w:t></w:r></w:p>")
    baseline, candidate = _pair(tmp_path, visible, hidden)
    changes = _by_attribute(diff(baseline, candidate))
    assert changes["format"] == {"p2"}
    assert "text" not in changes


def test_bookmark_removal_is_compared(tmp_path):
    with_mark = ('<w:p><w:bookmarkStart w:id="1" w:name="clause7"/>'
                 '<w:bookmarkEnd w:id="1"/><w:r><w:t>Clause 7</w:t></w:r></w:p>')
    without = "<w:p><w:r><w:t>Clause 7</w:t></w:r></w:p>"
    baseline, candidate = _pair(tmp_path, with_mark, without)
    assert _by_attribute(diff(baseline, candidate))["bookmark"] == {"p2"}


def test_word_autosave_bookmark_is_ignored(tmp_path):
    """Word writes _GoBack on every save; it is not a document change."""
    without = "<w:p><w:r><w:t>Clause 7</w:t></w:r></w:p>"
    with_goback = ('<w:p><w:bookmarkStart w:id="1" w:name="_GoBack"/>'
                   '<w:bookmarkEnd w:id="1"/><w:r><w:t>Clause 7</w:t></w:r></w:p>')
    baseline, candidate = _pair(tmp_path, without, with_goback)
    assert diff(baseline, candidate) == []


# --- properties below the text ------------------------------------------

def _paragraph(properties):
    return f"<w:p><w:pPr>{properties}</w:pPr><w:r><w:t>Clause</w:t></w:r></w:p>"


def test_paragraph_formatting_is_compared(tmp_path):
    baseline, candidate = _pair(
        tmp_path,
        _paragraph('<w:jc w:val="left"/><w:ind w:left="0"/>'),
        _paragraph('<w:jc w:val="center"/><w:ind w:left="2880"/>'),
    )
    changes = _by_attribute(diff(baseline, candidate))
    assert changes["paragraph_format"] == {"p2"}
    assert "text" not in changes


def test_tab_stops_are_compared(tmp_path):
    baseline, candidate = _pair(
        tmp_path,
        _paragraph('<w:tabs><w:tab w:val="right" w:pos="7920"/></w:tabs>'),
        _paragraph('<w:tabs><w:tab w:val="right" w:pos="1728"/></w:tabs>'),
    )
    assert _by_attribute(diff(baseline, candidate))["paragraph_format"] == {"p2"}


def test_list_level_change_is_compared(tmp_path):
    baseline, candidate = _pair(
        tmp_path,
        _paragraph('<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'),
        _paragraph('<w:numPr><w:ilvl w:val="3"/><w:numId w:val="1"/></w:numPr>'),
    )
    assert _by_attribute(diff(baseline, candidate))["paragraph_format"] == {"p2"}


def _table(cell_properties):
    return (
        '<w:tbl><w:tblPr><w:tblBorders><w:top w:val="single" w:sz="4"/>'
        "</w:tblBorders></w:tblPr>"
        "<w:tblGrid><w:gridCol w:w=\"4000\"/></w:tblGrid>"
        f"<w:tr><w:tc><w:tcPr>{cell_properties}</w:tcPr>"
        "<w:p><w:r><w:t>cell</w:t></w:r></w:p></w:tc></w:tr></w:tbl>"
    )


def test_table_cell_shading_and_width_are_compared(tmp_path):
    baseline, candidate = _pair(
        tmp_path,
        _table('<w:tcW w:w="4000" w:type="dxa"/>'),
        _table('<w:tcW w:w="200" w:type="dxa"/><w:shd w:val="clear" w:fill="000000"/>'),
    )
    changes = _by_attribute(diff(baseline, candidate))
    assert changes["table"] == {"table1!r1c1"}
    assert "text" not in changes


def test_table_borders_are_compared(tmp_path):
    bordered = _table("")
    borderless = bordered.replace('<w:tblBorders><w:top w:val="single" w:sz="4"/>'
                                  "</w:tblBorders>", "")
    baseline, candidate = _pair(tmp_path, bordered, borderless)
    assert _by_attribute(diff(baseline, candidate))["table"] == {"table1"}


# --- displayed images ----------------------------------------------------

def test_picture_removed_from_the_page_is_caught(fixtures, tmp_path):
    """The media part stays in the zip; the picture is gone from the page."""
    from generate import rewrite_zip

    source = fixtures["word_image_baseline"]
    with zipfile.ZipFile(source) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    start = xml.find("<w:drawing>")
    end = xml.find("</w:drawing>") + len("</w:drawing>")
    assert start != -1, "fixture should contain a drawing"

    baseline = tmp_path / "base.docx"
    candidate = tmp_path / "cand.docx"
    rewrite_zip(source, baseline, add={})
    rewrite_zip(baseline, candidate,
                add={"word/document.xml": (xml[:start] + xml[end:]).encode("utf-8")})

    with zipfile.ZipFile(candidate) as zf:
        assert any(n.startswith("word/media/") for n in zf.namelist()), \
            "the media part must still be there"

    changes = _by_attribute(diff(baseline, candidate))
    assert "images" in changes
    assert not check(baseline, candidate, STRICT).passed


def test_snapshot_reports_displayed_images_only(fixtures):
    assert snapshot(fixtures["word_image_baseline"])["images"]
    assert snapshot(fixtures["word_baseline"])["images"] == []


# --- the locations that already existed keep meaning what they meant ----

def test_body_paragraph_numbering_is_unchanged_by_nested_containers(tmp_path):
    """A content control between paragraphs must not renumber the body."""
    body = ("<w:p><w:r><w:t>one</w:t></w:r></w:p>"
            + _sdt("boxed")
            + "<w:p><w:r><w:t>two</w:t></w:r></w:p>")
    document = _build(tmp_path, "doc", body)
    paragraphs = snapshot(document)["paragraphs"]
    assert [p["text"] for p in paragraphs] == ["Intro paragraph.", "one", "two"]


def test_allowed_body_edit_is_reported_once(tmp_path):
    """An edit to body text must not also surface as a container change."""
    body = _sdt("250,000") + "<w:p><w:r><w:t>Original sentence.</w:t></w:r></w:p>"
    changed = _sdt("250,000") + "<w:p><w:r><w:t>Rewritten sentence.</w:t></w:r></w:p>"
    baseline, candidate = _pair(tmp_path, body, changed)
    changes = diff(baseline, candidate)
    assert [(c.location, c.attribute) for c in changes] == [("p2", "text")]
    assert check(baseline, candidate, BODY_TEXT_ONLY).passed
