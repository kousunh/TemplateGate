"""The false-positive guardrail for the whole snapshot surface.

Two properties, checked against every fixture:

1. Re-saving a document with the editing library and changing nothing must
   report nothing.  Any change here is noise the gate would cry wolf about.
2. Re-saving with exactly one allowed edit must report exactly that edit.
   This is the one that catches a package-layer part being hashed when its
   contents are already compared semantically — editing a cell rewrites the
   worksheet part by definition.

Both are what makes default-deny over the whole package safe.
"""

import pytest
from docx import Document
from openpyxl import load_workbook

from templategate import diff

# The *_package_* fixtures are left out on purpose: they carry parts that the
# editing libraries genuinely destroy on save, which is the opposite property
# and is asserted separately below.
EXCEL_FIXTURES = [
    "excel_baseline", "excel_good", "excel_bad",
    "excel_image_baseline", "excel_image_removed", "excel_native_rich",
]
WORD_FIXTURES = [
    "word_baseline", "word_good", "word_bad", "word_reformatted",
    "word_image_baseline", "word_image_removed",
]


@pytest.mark.parametrize("name", EXCEL_FIXTURES)
def test_openpyxl_roundtrip_reports_nothing(fixtures, tmp_path, name):
    source = fixtures[name]
    resaved = tmp_path / f"{name}.xlsx"
    load_workbook(source).save(resaved)
    assert diff(source, resaved) == []


@pytest.mark.parametrize("name", WORD_FIXTURES)
def test_python_docx_roundtrip_reports_nothing(fixtures, tmp_path, name):
    source = fixtures[name]
    resaved = tmp_path / f"{name}.docx"
    Document(str(source)).save(str(resaved))
    assert diff(source, resaved) == []


def test_openpyxl_roundtrip_with_one_edit_reports_only_that_edit(fixtures, tmp_path):
    """Editing a cell rewrites the worksheet part; only the cell may be reported."""
    source = fixtures["excel_native_rich"]
    edited = tmp_path / "edited.xlsx"
    workbook = load_workbook(source)
    workbook["Data"]["B2"] = 999
    workbook.save(edited)
    assert [(c.location, c.attribute) for c in diff(source, edited)] == [
        ("Data!B2", "value")
    ]


def test_python_docx_roundtrip_with_one_edit_reports_only_that_edit(fixtures, tmp_path):
    """Editing a paragraph rewrites document.xml; only the paragraph may be reported."""
    source = fixtures["word_baseline"]
    edited = tmp_path / "edited.docx"
    document = Document(str(source))
    document.paragraphs[3].text = "書き換えられた本文。"
    document.save(str(edited))
    assert [(c.location, c.attribute) for c in diff(source, edited)] == [
        ("p4", "text")
    ]


def test_library_roundtrip_reports_the_parts_the_library_destroys(fixtures, tmp_path):
    """The product's whole point: a save that silently drops parts must FAIL.

    Re-saving through openpyxl discards the chart, pivot tables, VBA project,
    embedded object and custom XML, none of which it can model.
    """
    source = fixtures["excel_package_baseline"]
    resaved = tmp_path / "resaved.xlsx"
    load_workbook(source).save(resaved)
    lost = {c.attribute for c in diff(source, resaved) if c.new is None}
    assert {"vba", "charts", "pivot_tables", "embedded", "custom_xml"} <= lost
